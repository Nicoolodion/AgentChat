"""
Chatinterface Agent Sandbox HTTP API Server
Provides isolated code execution, file operations, and document conversion.
"""

import argparse
import base64
import fcntl
import io
import json
import logging
import mimetypes
import os
import queue
import re
import shutil
import signal
import subprocess
import sys
import threading
import time
import traceback
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Optional

from flask import Flask, request, jsonify, Response

# Local helper module (same directory): persistent, streaming, leak-proof
# Python execution. See python_exec.py for details.
from python_exec import run_python, build_sanitized_env
from isolation import (
    alloc_ids,
    as_session_uid,
    drop_to_session,
    prepare_session_with_migration,
    session_home,
)
import ocr as ocr_engine

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    stream=sys.stdout,
)
logger = logging.getLogger("sandbox")

app = Flask(__name__)

# Configuration
WORKSPACE_ROOT = Path(os.environ.get("SANDBOX_WORKSPACE", "/workspace"))
MAX_EXECUTION_TIME = int(os.environ.get("SANDBOX_MAX_EXECUTION_TIME", "300"))
DEFAULT_TIMEOUT = int(os.environ.get("SANDBOX_DEFAULT_TIMEOUT", "60"))
MAX_OUTPUT_SIZE = 10 * 1024 * 1024  # 10MB

# Per-session UID isolation (see isolation.py). The server runs as root and
# chowns each session directory to a dedicated uid (mode 0700). User code runs
# as that uid via a full privilege drop, so one session cannot touch another's
# files (audit 2.1/2.3).


def _run_as_alloc(
    session_id: str,
    argv: list[str],
    env: dict,
    cwd: str,
    timeout: int,
    on_chunk: Optional[Callable[[str, str], None]] = None,
) -> dict:
    """Run ``argv`` as the session's dedicated uid with a process-group timeout.

    Returns ``{stdout, stderr, exit_code, timed_out, error}``.
    """
    proc = subprocess.Popen(
        argv,
        cwd=cwd,
        env=env,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        bufsize=1,
        start_new_session=True,
        preexec_fn=drop_to_session(session_id),
    )

    out_q: "queue.Queue[tuple[str, str]]" = queue.Queue()

    def reader(pipe, tag: str) -> None:
        try:
            for line in pipe:
                out_q.put((tag, line))
        except Exception:
            pass
        finally:
            out_q.put((tag, ""))

    t_out = threading.Thread(target=reader, args=(proc.stdout, "stdout"), daemon=True)
    t_err = threading.Thread(target=reader, args=(proc.stderr, "stderr"), daemon=True)
    t_out.start()
    t_err.start()

    eof = {"stdout": False, "stderr": False}
    stdout_buf: list[str] = []
    stderr_buf: list[str] = []
    timed_out = False
    deadline = time.time() + timeout
    while True:
        remaining = deadline - time.time()
        if remaining <= 0:
            timed_out = True
            break
        try:
            tag, line = out_q.get(timeout=remaining)
        except queue.Empty:
            timed_out = True
            break
        if line == "":
            eof[tag] = True
            if all(eof.values()):
                break
            continue
        stdout_buf.append(line) if tag == "stdout" else stderr_buf.append(line)
        if on_chunk:
            try:
                on_chunk(tag, line)
            except Exception:
                pass

    if timed_out:
        try:
            os.killpg(os.getpgid(proc.pid), signal.SIGKILL)
        except Exception:
            try:
                proc.kill()
            except Exception:
                pass
        try:
            proc.wait(timeout=5)
        except Exception:
            pass
        while True:
            try:
                tag, line = out_q.get_nowait()
            except queue.Empty:
                break
            if line == "":
                continue
            stdout_buf.append(line) if tag == "stdout" else stderr_buf.append(line)
            if on_chunk:
                try:
                    on_chunk(tag, line)
                except Exception:
                    pass
        return {
            "stdout": "".join(stdout_buf)[-MAX_OUTPUT_SIZE:],
            "stderr": "".join(stderr_buf)[-MAX_OUTPUT_SIZE:],
            "exit_code": -1,
            "timed_out": True,
            "error": f"timed out after {timeout}s",
        }

    try:
        proc.wait(timeout=10)
    except Exception:
        pass
    while True:
        try:
            tag, line = out_q.get_nowait()
        except queue.Empty:
            break
        if line == "":
            continue
        stdout_buf.append(line) if tag == "stdout" else stderr_buf.append(line)
        if on_chunk:
            try:
                on_chunk(tag, line)
            except Exception:
                pass
    t_out.join(timeout=2)
    t_err.join(timeout=2)
    return {
        "stdout": "".join(stdout_buf)[-MAX_OUTPUT_SIZE:],
        "stderr": "".join(stderr_buf)[-MAX_OUTPUT_SIZE:],
        "exit_code": proc.returncode if proc.returncode is not None else -1,
        "timed_out": False,
        "error": None,
    }


# Ensure workspace exists
WORKSPACE_ROOT.mkdir(parents=True, exist_ok=True)


def make_error(message: str, status_code: int = 400) -> tuple[Response, int]:
    return jsonify({"error": message, "success": False}), status_code


def make_success(data: dict[str, Any]) -> Response:
    return jsonify({"success": True, **data})


# session_id values are interpolated into workspace paths and used as alloc
# uid keys. They must never contain path separators or traversal sequences.
# The TS client always sends UUID/cuid ids; "default" is tolerated for legacy
# callers. Every other value must match this safe pattern.
_VALID_SESSION_ID_RE = re.compile(r"^[A-Za-z0-9_-]{8,128}$")


@app.before_request
def _validate_session_id() -> Optional[tuple]:
    if request.method != "POST":
        return None
    data = request.get_json(silent=True)
    if not isinstance(data, dict) or "session_id" not in data:
        return None
    sid = data.get("session_id")
    if sid == "default":
        return None
    if not isinstance(sid, str) or not _VALID_SESSION_ID_RE.match(sid):
        return make_error(
            "Invalid session_id; must match ^[A-Za-z0-9_-]{8,128}$", 400
        )
    return None


def _capture_failure(prefix: str, jr: dict) -> tuple[Response, int]:
    """Format an alloc-subprocess failure/timeout including captured output.

    ``_run_as_alloc`` already captures stdout/stderr even on timeout; surfacing
    them (mirroring ``_build_pptx_result``) makes opaque "timed out" / failure
    messages self-explanatory in the agent UI.
    """
    stdout = (jr.get("stdout") or "").strip()
    stderr = (jr.get("stderr") or "").strip()
    parts = [prefix]
    if stdout:
        parts.append(f"stdout:\n{stdout}")
    if stderr:
        parts.append(f"stderr:\n{stderr}")
    status = 504 if jr.get("timed_out") else 500
    return make_error("\n".join(parts).strip(), status)


def resolve_workspace_path(session_id: str, sub_path: str = "") -> Path:
    """Resolve a path within the session workspace, preventing traversal."""
    session_workspace = (WORKSPACE_ROOT / session_id).resolve()
    session_workspace.mkdir(parents=True, exist_ok=True)

    if sub_path:
        target = (session_workspace / sub_path.lstrip("/")).resolve()
    else:
        target = session_workspace

    # Security: prevent path traversal outside the workspace. relative_to
    # raises ValueError if target is not within session_workspace, which is
    # stricter and symlink-safer than a str().startswith() prefix check.
    try:
        target.relative_to(session_workspace)
    except ValueError:
        raise ValueError("Path traversal detected")

    return target


def _chown_alloc(path: Path, session_id: str) -> None:
    """Best-effort: hand ownership of a server-created path to the session uid."""
    auid, agid = alloc_ids(session_id)
    try:
        os.chown(path, auid, agid)
    except OSError:
        pass


def _session_home(session_id: str) -> Path:
    """Per-session alloc-owned HOME (see :func:`isolation.session_home`)."""
    return session_home(session_id)


def _alloc_run_env(session_id: str, *, base: Optional[dict] = None) -> dict:
    """Sanitized env for a subprocess that will drop to the session's alloc uid.

    HOME and all XDG/NUGET profile/cache vars point at the per-session .home so
    the dropped process never reads or writes the shared /tmp cache (which would
    either collide with or be blocked by root-owned entries). TMPDIR stays
    /tmp (world-writable tmpfs, mode 1777) for scratch files.
    """
    home = _session_home(session_id)
    env = build_sanitized_env(base)
    env["HOME"] = str(home)
    env["XDG_CACHE_HOME"] = str(home / ".cache")
    env["XDG_CONFIG_HOME"] = str(home / ".config")
    env["XDG_DATA_HOME"] = str(home / ".local" / "share")
    env["DOTNET_CLI_HOME"] = str(home)
    env["NUGET_PACKAGES"] = str(home / "nuget" / "packages")
    env["NUGET_HTTP_CACHE_PATH"] = str(home / "nuget" / "http-cache")
    env["NUGET_SCRATCH"] = str(home / "nuget" / "scratch")
    env["TMPDIR"] = "/tmp"
    return env


def ensure_session_dirs(session_id: str) -> Path:
    """Create the session directory structure owned by its dedicated uid.

    Runs as root: chowns the tree to the session's ``alloc`` uid (mode 0700),
    so user code (which drops to that uid) can access its own files but no
    sibling session can. Subsequent file/docx ops switch effective uid to
    ``alloc`` to touch these files.
    """
    try:
        return prepare_session_with_migration(session_id)
    except Exception as e:
        logger.warning("[isolation] prepare error for %s: %s", session_id, e)
        session_workspace = WORKSPACE_ROOT / session_id
        session_workspace.mkdir(parents=True, exist_ok=True)
        for sub in ("upload", "output", "temp"):
            (session_workspace / sub).mkdir(parents=True, exist_ok=True)
        return session_workspace


# ═══════════════════════════════════════════════════════════════════════════════
# Health
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/health", methods=["GET"])
def health() -> Response:
    """Health check endpoint."""
    checks = {
        "python": True,
        "node": False,
        "playwright": False,
        "libreoffice": False,
        "dotnet": False,
        "isolation": os.geteuid() == 0,
    }

    try:
        subprocess.run(["node", "--version"], capture_output=True, check=True, timeout=5)
        checks["node"] = True
    except Exception:
        pass

    try:
        import playwright
        checks["playwright"] = True
    except ImportError:
        pass

    try:
        subprocess.run(["libreoffice", "--version"], capture_output=True, check=True, timeout=5)
        checks["libreoffice"] = True
    except Exception:
        pass

    try:
        subprocess.run(["dotnet", "--version"], capture_output=True, check=True, timeout=5)
        checks["dotnet"] = True
    except Exception:
        pass

    all_ok = all(checks.values())
    return jsonify({
        "status": "ok" if all_ok else "degraded",
        "checks": checks,
        "workspace_root": str(WORKSPACE_ROOT),
    })


# ═══════════════════════════════════════════════════════════════════════════════
# Python Execution (IPython-like, persistent per session_id)
# ═══════════════════════════════════════════════════════════════════════════════

# Python execution lives in python_exec.run_python (persistent state, real-time
# streaming, process-group timeouts). The routes below just adapt it to HTTP.


@app.route("/exec/python", methods=["POST"])
def exec_python() -> Response:
    """Execute Python code with persistent session state."""
    data = request.get_json(force=True) or {}
    code = data.get("code", "")
    session_id = data.get("session_id", "default")
    timeout = int(data.get("timeout", DEFAULT_TIMEOUT))

    if not code:
        return make_error("Missing 'code' field", 400)

    logger.info(f"[python] session={session_id} timeout={timeout}")

    result = run_python(code, session_id, timeout, WORKSPACE_ROOT)
    return jsonify(result)


@app.route("/exec/python/stream", methods=["POST"])
def exec_python_stream():
    """Stream Python execution output as it happens.

    Emits a JSON object per line:
      {"t":"stdout","s":"..."} / {"t":"stderr","s":"..."}  — live chunks
      {"t":"result", ...}                                   — final result record
    """
    data = request.get_json(force=True) or {}
    code = data.get("code", "")
    session_id = data.get("session_id", "default")
    timeout = int(data.get("timeout", DEFAULT_TIMEOUT))

    if not code:
        return make_error("Missing 'code' field", 400)

    logger.info(f"[python/stream] session={session_id} timeout={timeout}")

    def generate():
        # Run execution in a worker thread; forward live chunks (stdout/stderr
        # lines) through a queue as newline-delimited JSON, then emit the
        # final result record.
        import queue as _q
        import threading as _t

        out_q: "_q.Queue" = _q.Queue()

        def worker():
            def cb(stream, text):
                out_q.put({"t": stream, "s": text})
            try:
                res = run_python(code, session_id, timeout, WORKSPACE_ROOT, on_chunk=cb)
                out_q.put({"t": "result", "data": res})
            except Exception as e:  # pragma: no cover
                out_q.put({"t": "result", "data": {
                    "stdout": "", "stderr": "", "images": [],
                    "error": str(e), "execution_time_ms": 0,
                }})

        th = _t.Thread(target=worker, daemon=True)
        th.start()
        while True:
            item = out_q.get(timeout=timeout + 10)
            yield json.dumps(item) + "\n"
            if item.get("t") == "result":
                break

    return Response(generate(), mimetype="application/x-ndjson")


# ═══════════════════════════════════════════════════════════════════════════════
# Shell Execution
# ═══════════════════════════════════════════════════════════════════════════════

BLACKLISTED_SHELL_PATTERNS = [
    re.compile(r"rm\s+-rf\s+/", re.IGNORECASE),
    re.compile(r"rm\s+-rf\s+/\*", re.IGNORECASE),
    re.compile(r"mkfs\.", re.IGNORECASE),
    re.compile(r":\(\)\{\s*:\|:&\s*\};:", re.IGNORECASE),
    re.compile(r"dd\s+if=/dev/zero", re.IGNORECASE),
    re.compile(r">\s*/dev/sda", re.IGNORECASE),
    re.compile(r"curl\s+.*\|\s*sh", re.IGNORECASE),
    re.compile(r"wget\s+.*\|\s*sh", re.IGNORECASE),
    re.compile(r"base64\s+.*\|\s*sh", re.IGNORECASE),
    re.compile(r"printf\s+.*\|\s*sh", re.IGNORECASE),
    re.compile(r"xargs\s+sh", re.IGNORECASE),
    re.compile(r"eval\s+", re.IGNORECASE),
    re.compile(r"exec\s+sh", re.IGNORECASE),
    re.compile(r"\$\(\s*rm\s", re.IGNORECASE),
    re.compile(r"chmod\s+[0-7]*777", re.IGNORECASE),
    re.compile(r"chown\s+.*root", re.IGNORECASE),
    re.compile(r"/etc/passwd", re.IGNORECASE),
    re.compile(r"/etc/shadow", re.IGNORECASE),
    # Shell is the one path that bypasses the file-tool path validator (audit
    # 2.3). Block obvious attempts to read/patch the API server, read other
    # sessions' data, or introspect the host/processes. This is defense-in-depth,
    # not airtight: a determined caller can evade string matching, which is why
    # OS-level execution containment (separate container / bwrap / Landlock) is
    # the real fix for findings 2.1 and 2.3.
    re.compile(r"/app/(lib|scripts|skills)", re.IGNORECASE),
    re.compile(r"/proc/(self|\d+)/(environ|fd|maps|mem)", re.IGNORECASE),
    re.compile(r"/sys/(kernel|class|devices)", re.IGNORECASE),
    re.compile(r"/\.dockerenv", re.IGNORECASE),
    re.compile(r">+\s*/app\b", re.IGNORECASE),
]


# A conservative PATH for session shell commands: user toolchain locations only.
# No write access to /app, /etc, /usr/local/bin etc. is implied by PATH; this
# just avoids leaking odd host binaries into the session env.
RESTRICTED_PATH = "/usr/local/bin:/usr/bin:/bin:/usr/sbin:/sbin"


def is_shell_blacklisted(command: str) -> bool:
    for pattern in BLACKLISTED_SHELL_PATTERNS:
        if pattern.search(command):
            return True
    return False


@app.route("/exec/shell", methods=["POST"])
def exec_shell() -> Response:
    """Execute a shell command in the session workspace."""
    data = request.get_json(force=True) or {}
    command = data.get("command", "")
    session_id = data.get("session_id", "default")
    working_dir = data.get("working_dir", "")
    timeout = int(data.get("timeout", DEFAULT_TIMEOUT))

    if not command:
        return make_error("Missing 'command' field", 400)

    if is_shell_blacklisted(command):
        return make_error("Command blocked by security policy", 403)

    session_workspace = ensure_session_dirs(session_id)
    cwd = session_workspace
    if working_dir:
        try:
            cwd = resolve_workspace_path(session_id, working_dir)
        except ValueError:
            return make_error("Invalid working directory", 400)

    logger.info(f"[shell] session={session_id} cmd={command[:80]}")

    # Run the command with a secret-scrubbed environment so credentials that
    # happen to be present in the container env are not handed to user code
    # (audit 2.5). HOME points at a per-session alloc-owned .home so tools the
    # agent shells out to (libreoffice, etc.) can write caches without polluting
    # or being blocked by root-owned /tmp entries. umask 077 keeps any files the
    # command creates private within the session's own uid (defense-in-depth for 2.1).
    shell_env = _alloc_run_env(session_id, base=os.environ.copy())
    shell_env["WORKSPACE_DIR"] = str(session_workspace)
    shell_env["PATH"] = RESTRICTED_PATH + os.pathsep + shell_env.get("PATH", "")

    start = time.time()

    # Run under the session's dedicated uid: it cannot reach any other
    # session's workspace (audit 2.1/2.3).
    try:
        jr = _run_as_alloc(
            session_id,
            ["/bin/sh", "-c", command],
            shell_env,
            str(cwd),
            min(timeout, MAX_EXECUTION_TIME),
        )
        duration_ms = int((time.time() - start) * 1000)
        return jsonify({
            "stdout": jr["stdout"],
            "stderr": jr["stderr"],
            "exit_code": jr["exit_code"],
            "error": jr.get("error"),
            "duration_ms": duration_ms,
        })
    except Exception as e:
        logger.exception("[shell] exec failed")
        return jsonify({
            "stdout": "",
            "stderr": str(e),
            "exit_code": -1,
            "error": str(e),
            "duration_ms": int((time.time() - start) * 1000),
        })


# ═══════════════════════════════════════════════════════════════════════════════
# File Operations
# ═══════════════════════════════════════════════════════════════════════════════

SKILLS_ROOT = Path(os.environ.get("SANDBOX_SKILLS_ROOT", "/app/skills"))


@app.route("/file/read", methods=["POST"])
def file_read() -> Response:
    data = request.get_json(force=True) or {}
    file_path = data.get("path", "")
    encoding = data.get("encoding", "utf8")
    session_id = data.get("session_id", "default")

    if not file_path:
        return make_error("Missing 'path' field", 400)

    # Allow reading skill files from /app/skills/ (mounted read-only)
    if file_path.startswith("/app/skills/") or file_path.startswith("skills/"):
        skill_rel = file_path.removeprefix("/app/skills/").removeprefix("skills/")
        target = (SKILLS_ROOT / skill_rel).resolve()
        if not str(target).startswith(str(SKILLS_ROOT.resolve())):
            return make_error("Path traversal detected", 400)
        if not target.exists():
            return make_error("File not found", 404)
        if target.is_dir():
            return make_error("Path is a directory", 400)
        try:
            if encoding == "base64":
                content = base64.b64encode(target.read_bytes()).decode("ascii")
            else:
                content = target.read_text(encoding="utf-8", errors="replace")
            return make_success({
                "content": content,
                "encoding": encoding,
                "size": target.stat().st_size,
                "modified_at": datetime.fromtimestamp(target.stat().st_mtime).isoformat(),
            })
        except Exception as e:
            return make_error(f"Read failed: {e}", 500)

    ensure_session_dirs(session_id)
    try:
        with as_session_uid(session_id):
            target = resolve_workspace_path(session_id, file_path)
            if not target.exists():
                return make_error("File not found", 404)
            if target.is_dir():
                return make_error("Path is a directory", 400)

            if encoding == "base64":
                content = base64.b64encode(target.read_bytes()).decode("ascii")
            else:
                content = target.read_text(encoding="utf-8", errors="replace")

            return make_success({
                "content": content,
                "encoding": encoding,
                "size": target.stat().st_size,
                "modified_at": datetime.fromtimestamp(target.stat().st_mtime).isoformat(),
            })
    except ValueError:
        return make_error("Invalid path", 400)
    except Exception as e:
        return make_error(f"Read failed: {e}", 500)


@app.route("/file/write", methods=["POST"])
def file_write() -> Response:
    data = request.get_json(force=True) or {}
    file_path = data.get("path", "")
    content = data.get("content", "")
    encoding = data.get("encoding", "utf8")
    session_id = data.get("session_id", "default")

    if not file_path:
        return make_error("Missing 'path' field", 400)

    ensure_session_dirs(session_id)
    try:
        with as_session_uid(session_id):
            target = resolve_workspace_path(session_id, file_path)
            target.parent.mkdir(parents=True, exist_ok=True)
            if encoding == "base64":
                target.write_bytes(base64.b64decode(content))
            else:
                target.write_text(content, encoding="utf-8")

            return make_success({
                "path": str(target),
                "size": target.stat().st_size,
            })
    except ValueError:
        return make_error("Invalid path", 400)
    except Exception as e:
        return make_error(f"Write failed: {e}", 500)


@app.route("/file/list", methods=["POST"])
def file_list() -> Response:
    data = request.get_json(force=True) or {}
    dir_path = data.get("path", "/")
    session_id = data.get("session_id", "default")

    # Allow listing skill files from /app/skills/
    if dir_path.startswith("/app/skills/") or dir_path.startswith("skills/"):
        skill_rel = dir_path.removeprefix("/app/skills/").removeprefix("skills/")
        target = (SKILLS_ROOT / skill_rel).resolve() if skill_rel else SKILLS_ROOT.resolve()
        if not str(target).startswith(str(SKILLS_ROOT.resolve())):
            return make_error("Path traversal detected", 400)
        if not target.exists():
            return make_error("Directory not found", 404)
        if not target.is_dir():
            return make_error("Path is not a directory", 400)
        files = []
        try:
            for entry in target.iterdir():
                stat_info = entry.stat()
                mime_type, _ = mimetypes.guess_type(str(entry)) if entry.is_file() else (None, None)
                files.append({
                    "name": entry.name,
                    "path": str(entry),
                    "is_directory": entry.is_dir(),
                    "size": stat_info.st_size if entry.is_file() else 0,
                    "mime_type": mime_type,
                    "modified_at": datetime.fromtimestamp(stat_info.st_mtime).isoformat(),
                })
            files.sort(key=lambda x: (not x["is_directory"], x["name"].lower()))
            return make_success({"files": files})
        except Exception as e:
            return make_error(f"List failed: {e}", 500)

    ensure_session_dirs(session_id)
    try:
        with as_session_uid(session_id):
            target = resolve_workspace_path(session_id, dir_path)
            if not target.exists():
                return make_error("Directory not found", 404)
            if not target.is_dir():
                return make_error("Path is not a directory", 400)

            files = []
            for entry in target.iterdir():
                stat_info = entry.stat()
                mime_type, _ = mimetypes.guess_type(str(entry)) if entry.is_file() else (None, None)
                files.append({
                    "name": entry.name,
                    "path": str(entry),
                    "is_directory": entry.is_dir(),
                    "size": stat_info.st_size if entry.is_file() else 0,
                    "mime_type": mime_type,
                    "modified_at": datetime.fromtimestamp(stat_info.st_mtime).isoformat(),
                })

            files.sort(key=lambda x: (not x["is_directory"], x["name"].lower()))
            return make_success({"files": files})
    except ValueError:
        return make_error("Invalid path", 400)
    except Exception as e:
        return make_error(f"List failed: {e}", 500)


@app.route("/file/delete", methods=["POST"])
def file_delete() -> Response:
    data = request.get_json(force=True) or {}
    file_path = data.get("path", "")
    session_id = data.get("session_id", "default")

    if not file_path:
        return make_error("Missing 'path' field", 400)

    ensure_session_dirs(session_id)
    try:
        with as_session_uid(session_id):
            target = resolve_workspace_path(session_id, file_path)
            if not target.exists():
                return make_error("File not found", 404)

            if target.is_dir():
                shutil.rmtree(target)
            else:
                target.unlink()
            return make_success({"deleted": str(target)})
    except ValueError:
        return make_error("Invalid path", 400)
    except Exception as e:
        return make_error(f"Delete failed: {e}", 500)


@app.route("/file/move", methods=["POST"])
def file_move() -> Response:
    data = request.get_json(force=True) or {}
    source = data.get("source", "")
    destination = data.get("destination", "")
    session_id = data.get("session_id", "default")

    if not source or not destination:
        return make_error("Missing 'source' or 'destination' field", 400)

    ensure_session_dirs(session_id)
    try:
        with as_session_uid(session_id):
            src = resolve_workspace_path(session_id, source)
            dst = resolve_workspace_path(session_id, destination)
            if not src.exists():
                return make_error("Source not found", 404)

            dst.parent.mkdir(parents=True, exist_ok=True)
            src.rename(dst)
            return make_success({"source": str(src), "destination": str(dst)})
    except ValueError:
        return make_error("Invalid path", 400)
    except Exception as e:
        return make_error(f"Move failed: {e}", 500)


@app.route("/file/info", methods=["POST"])
def file_info() -> Response:
    data = request.get_json(force=True) or {}
    file_path = data.get("path", "")
    session_id = data.get("session_id", "default")

    if not file_path:
        return make_error("Missing 'path' field", 400)

    ensure_session_dirs(session_id)
    try:
        with as_session_uid(session_id):
            target = resolve_workspace_path(session_id, file_path)
            if not target.exists():
                return make_error("File not found", 404)

            stat_info = target.stat()
            mime_type, _ = mimetypes.guess_type(str(target)) if target.is_file() else (None, None)
            return make_success({
                "name": target.name,
                "path": str(target),
                "size": stat_info.st_size if target.is_file() else 0,
                "mime_type": mime_type,
                "modified_at": datetime.fromtimestamp(stat_info.st_mtime).isoformat(),
                "is_directory": target.is_dir(),
            })
    except ValueError:
        return make_error("Invalid path", 400)
    except Exception as e:
        return make_error(f"Info failed: {e}", 500)


# ═══════════════════════════════════════════════════════════════════════════════
# Web rendering (Playwright headless) — for JS-heavy SPAs whose static HTML
# is an empty shell. Used as a fallback by the agent's web_fetch tool.
# ═══════════════════════════════════════════════════════════════════════════════

def _block_private_route(url: str) -> bool:
    """Reject private/loopback/internal hosts before launching a browser at them."""
    from urllib.parse import urlparse
    host = (urlparse(url).hostname or "").lower().strip("[]")
    blocked = {"localhost", "metadata.google.internal", "metadata",
              "169.254.169.254", "metadata.aws.internal"}
    if host in blocked:
        return True
    if host.endswith(".internal") or host.endswith(".local") or host.endswith(".localhost"):
        return True
    import re as _re
    v4 = _re.match(r"^(\d{1,3})\.(\d{1,3})\.(\d{1,3})\.(\d{1,3})$", host)
    if v4:
        a, b = int(v4.group(1)), int(v4.group(2))
        if a in (0, 10, 127) or (a == 169 and b == 254) or (a == 172 and 16 <= b <= 31) \
                or (a == 192 and b == 168) or (a == 100 and 64 <= b <= 127):
            return True
    if host in ("::1", "::") or host.startswith("fe80:") or host.startswith("fc") or host.startswith("fd"):
        return True
    return False


@app.route("/web/render", methods=["POST"])
def web_render() -> Response:
    """Render a URL with a headless Chromium and return the fully-built HTML.

    Optional:
      - cookies: list of {name, value, domain?, path?} for authenticated
        browsing of user-specified sites.
      - wait_for: a CSS selector to wait for before capturing (default: waits
        for networkidle, capped, then a short timeout).
      - timeout: total seconds (default 35).
    """
    data = request.get_json(force=True) or {}
    url = data.get("url", "")
    cookies = data.get("cookies", []) or []
    wait_for = data.get("wait_for", "")
    timeout = min(int(data.get("timeout", 35)), 90)

    if not url:
        return make_error("Missing 'url' field", 400)
    if not (url.startswith("http://") or url.startswith("https://")):
        return make_error("Only http(s) URLs can be rendered", 400)
    if _block_private_route(url):
        return make_error("Blocked private/internal host", 400)

    logger.info(f"[web/render] url={url[:80]} cookies={len(cookies)} wait_for={wait_for}")

    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        return make_error("Playwright is not installed in the sandbox", 500)

    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True, args=["--no-sandbox"])
            context = browser.new_context(
                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                           "(KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36",
                viewport={"width": 1280, "height": 1800},
            )
            if cookies:
                norm_cookies = []
                for c in cookies:
                    if not isinstance(c, dict) or not c.get("name") or c.get("value") is None:
                        continue
                    norm_cookies.append({
                        "name": str(c["name"]),
                        "value": str(c["value"]),
                        "domain": c.get("domain") or None,
                        "path": c.get("path") or "/",
                        "httpOnly": bool(c.get("http_only", False)),
                        "secure": bool(c.get("secure", False)),
                    })
                if norm_cookies:
                    try:
                        context.add_cookies(norm_cookies)
                    except Exception as e:
                        logger.warning(f"[web/render] add_cookies failed: {e}")
            page = context.new_page()
            try:
                page.goto(url, wait_until="domcontentloaded", timeout=timeout * 1000)
            except Exception as e:
                browser.close()
                return make_error(f"Navigation failed: {e}", 502)

            try:
                if wait_for:
                    page.wait_for_selector(wait_for, timeout=timeout * 1000)
                else:
                    # Give SPAs time to hydrate / fetch data.
                    try:
                        page.wait_for_load_state("networkidle", timeout=min(timeout, 8) * 1000)
                    except Exception:
                        pass
            except Exception:
                pass

            html = page.content()
            final_url = page.url
            title = page.title()
            content_type = "text/html"
            browser.close()

        return make_success({
            "url": url,
            "final_url": final_url,
            "title": title,
            "content_type": content_type,
            "html": html[:2_000_000],
            "size": len(html),
        })
    except Exception as e:
        return make_error(f"Render failed: {e}", 500)


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX Read (extract structured content from .docx files)
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/docx/read", methods=["POST"])
def docx_read() -> Response:
    """Parse a .docx file and return structured content (paragraphs, tables, images, styles)."""
    data = request.get_json(force=True) or {}
    file_path = data.get("path", "")
    session_id = data.get("session_id", "default")
    include_images = data.get("include_images", True)

    if not file_path:
        return make_error("Missing 'path' field", 400)

    ensure_session_dirs(session_id)
    try:
        with as_session_uid(session_id):
            target = resolve_workspace_path(session_id, file_path)
            if not target.exists():
                return make_error(f"File not found: {file_path}", 404)
            if not str(target).lower().endswith((".docx", ".doc")):
                return make_error("File must be a .docx or .doc file", 400)
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn

            doc = Document(str(target))

            paragraphs = []
        tables = []
        images = []
        image_idx = 0

        # Extract paragraphs with style info
        for i, p in enumerate(doc.paragraphs):
            style_name = p.style.name if p.style else ""
            alignment = ""
            try:
                if p.alignment is not None:
                    alignment_map = {
                        WD_ALIGN_PARAGRAPH.LEFT: "left",
                        WD_ALIGN_PARAGRAPH.CENTER: "center",
                        WD_ALIGN_PARAGRAPH.RIGHT: "right",
                        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
                    }
                    alignment = alignment_map.get(p.alignment, "")
            except Exception:
                jc_map = {"left": "left", "start": "left", "center": "center",
                          "right": "right", "end": "right", "justify": "justify",
                          "both": "justify", "distribute": "justify"}
                try:
                    pPr = p._element.find(qn('w:pPr'))
                    if pPr is not None:
                        jc = pPr.find(qn('w:jc'))
                        if jc is not None:
                            val = jc.get(qn('w:val'), '')
                            alignment = jc_map.get(val, val)
                except Exception:
                    alignment = ""
            
            is_heading = style_name.startswith("Heading")
            heading_level = int(style_name.replace("Heading ", "").replace("Heading", "1") or "0") if is_heading else 0
            
            text = p.text.strip()
            
            # Check for inline images in this paragraph
            has_image = False
            for run in p.runs:
                if run._element.xml.find("blip") != -1:
                    has_image = True
                    break
            
            if text or has_image:
                entry = {
                    "index": i,
                    "text": text,
                    "style": style_name,
                    "alignment": alignment,
                    "is_heading": is_heading,
                    "heading_level": heading_level,
                    "has_image": has_image,
                }
                paragraphs.append(entry)

        # Extract tables
        for t_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            tables.append({
                "index": t_idx,
                "rows": len(rows),
                "columns": len(rows[0]) if rows else 0,
                "data": rows,
            })

        # Extract images
        if include_images:
            try:
                from docx.opc.constants import RELATIONSHIP_TYPE as RT
                import base64
                
                for rel in doc.part.rels.values():
                    if "image" in rel.reltype:
                        try:
                            image_data = rel.target_part.blob
                            ext = rel.target_part.partname.split(".")[-1] if "." in str(rel.target_part.partname) else "png"
                            mime_map = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg", "gif": "image/gif", "bmp": "image/bmp", "webp": "image/webp", "tiff": "image/tiff", "emf": "image/x-emf", "wmf": "image/x-wmf"}
                            mime = mime_map.get(ext.lower(), "image/png")
                            
                            # Skip EMF/WMF as they can't be displayed in browser
                            if ext.lower() in ("emf", "wmf"):
                                images.append({
                                    "index": image_idx,
                                    "mime_type": mime,
                                    "size": len(image_data),
                                    "extension": ext,
                                    "note": "EMF/WMF format - not displayable as base64",
                                })
                                image_idx += 1
                                continue
                            
                            b64 = base64.b64encode(image_data).decode("ascii")
                            images.append({
                                "index": image_idx,
                                "mime_type": mime,
                                "size": len(image_data),
                                "extension": ext,
                                "data_url": f"data:{mime};base64,{b64[:100]}...",
                            })
                            image_idx += 1
                        except Exception:
                            pass
            except Exception:
                pass

        # Build a readable text summary
        text_parts = []
        for p in paragraphs:
            prefix = "#" * p["heading_level"] + " " if p["is_heading"] else ""
            text_parts.append(f"{prefix}{p['text']}")
        for t in tables:
            text_parts.append(f"\n[Table {t['index']}: {t['rows']}x{t['columns']}]")
            for row in t["data"]:
                text_parts.append(" | ".join(row))
        text_summary = "\n".join(text_parts)

        return make_success({
            "path": file_path,
            "paragraphs": paragraphs,
            "tables": tables,
            "images": images,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "image_count": len(images),
            "text_summary": text_summary,
        })

    except ValueError:
        return make_error("Invalid path", 400)
    except Exception as e:
        return make_error(f"Failed to parse .docx: {e}", 500)


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX Template Fill (high-level: preserve template layout, replace body content)
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/docx/template-fill", methods=["POST"])
def docx_template_fill() -> Response:
    """
    Fill a .docx template with new content while preserving the template's
    cover page, tables, headers/footers, and styles.
    
    The agent provides:
    - template_path: path to the template .docx
    - output_path: where to save the result
    - sections: list of {heading, content, images[]} to replace body content
    - keep_cover_page: whether to preserve the first page/table(s)
    - cover_replacements: optional {search: replace} for text in the cover page
    
    This eliminates the need for window-by-window WIR editing when the task is
    "follow this template as an example and create a new document."
    """
    data = request.get_json(force=True) or {}
    session_id = data.get("session_id", "default")
    template_path = data.get("template_path", "")
    output_path = data.get("output_path", "")
    sections = data.get("sections", [])
    keep_cover_page = data.get("keep_cover_page", True)
    cover_replacements = data.get("cover_replacements", {})
    
    if not template_path:
        return make_error("Missing 'template_path' field", 400)
    if not output_path:
        return make_error("Missing 'output_path' field", 400)
    if not sections:
        return make_error("Missing 'sections' field — provide at least one section", 400)
    
    session_workspace = ensure_session_dirs(session_id)

    try:
        template_file = resolve_workspace_path(session_id, template_path)
    except ValueError:
        return make_error("Invalid template_path", 400)

    if not template_file.exists():
        return make_error(f"Template file not found: {template_path}", 404)

    try:
        out_file = resolve_workspace_path(session_id, output_path)
    except ValueError:
        return make_error("Invalid output_path", 400)

    out_file.parent.mkdir(parents=True, exist_ok=True)
    _chown_alloc(out_file.parent, session_id)
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import parse_xml as docx_oxml_parse_xml
        docx_Pt = Pt
        import copy
        import re
        
        # Read the template to extract style definitions
        template_doc = Document(str(template_file))
        
        # Create a new document based on the template (this preserves styles)
        # We copy the template then selectively remove body content
        import shutil
        shutil.copy2(str(template_file), str(out_file))
        doc = Document(str(out_file))
        
        # Apply cover page text replacements
        if cover_replacements:
            for search_text, replace_text in cover_replacements.items():
                seen_paragraphs = set()
                for paragraph in doc.paragraphs:
                    if search_text in paragraph.text:
                        pid = id(paragraph._element)
                        if pid not in seen_paragraphs:
                            seen_paragraphs.add(pid)
                            for run in paragraph.runs:
                                if search_text in run.text:
                                    run.text = run.text.replace(search_text, replace_text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if search_text in paragraph.text:
                                    pid = id(paragraph._element)
                                    if pid not in seen_paragraphs:
                                        seen_paragraphs.add(pid)
                                        for run in paragraph.runs:
                                            if search_text in run.text:
                                                run.text = run.text.replace(search_text, replace_text)
        
        # Determine where the body content starts.
        # Strategy: Find the first element that is NOT part of the cover page.
        # The cover page consists of: the first table (protocol header) and any
        # paragraphs WITHIN/BEFORE that table. Everything after the last cover-page
        # table element is body content and should be removed.
        cover_end_index = 0
        if keep_cover_page:
            body = doc.element.body
            table_elements = body.findall(qn('w:tbl'))
            
            cover_end_elem = None
            
            if table_elements:
                last_cover_table = table_elements[0]
                
                for child in body:
                    if child is last_cover_table:
                        cover_end_elem = child
                        break
                
                if cover_end_elem is not None:
                    found_table = False
                    to_remove = []
                    for child in body:
                        if child is cover_end_elem:
                            found_table = True
                            continue
                        if found_table:
                            to_remove.append(child)
                    
                    for elem in to_remove:
                        body.remove(elem)
        
        # Optional: insert an automatic Table of Contents (built from the
        # headings added below) right after the cover page. Word/LibreOffice
        # populate the page numbers + entry text on open ("update field").
        include_toc = data.get("include_toc", False)
        toc_paragraph = None
        if include_toc:
            toc_paragraph = doc.add_paragraph()
            run = toc_paragraph.add_run()
            # TOC field code: { TOC \o "1-3" \h \z \u } — levels 1-3, hyperlinks.
            fldChar_begin = docx_oxml_parse_xml(
                '<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            )
            instrText = docx_oxml_parse_xml(
                '<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'
            )
            fldChar_sep = docx_oxml_parse_xml(
                '<w:fldChar w:fldCharType="separate" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            )
            fldChar_end = docx_oxml_parse_xml(
                '<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            )
            run._element.append(fldChar_begin)
            run._element.append(instrText)
            run._element.append(fldChar_sep)
            run._element.append(fldChar_end)
            # Placeholder text so the TOC isn't invisible before first refresh.
            placeholder = doc.add_paragraph(
                "Right-click and choose “Update Field” to build the table of contents."
            )
            try:
                placeholder.runs[0].font.italic = True
                placeholder.runs[0].font.size = docx_Pt(9)
            except Exception:
                pass

        # Now add new content sections
        for section in sections:
            heading = section.get("heading", "")
            content = section.get("content", "")
            images = section.get("images", [])
            
            # Add heading
            if heading:
                heading_level = section.get("heading_level", 1)
                style_name = f"Heading {heading_level}"
                try:
                    h_para = doc.add_heading(heading, level=heading_level)
                except Exception:
                    h_para = doc.add_paragraph(heading)
                    h_para.style = doc.styles[style_name] if style_name in [s.name for s in doc.styles] else None
            
            # Add content (parse simple markdown-like formatting)
            if content:
                lines = content.split('\n')
                prev_was_blank = False
                for line in lines:
                    line_stripped = line.strip()
                    if not line_stripped:
                        if not prev_was_blank:
                            pass
                        prev_was_blank = True
                        continue
                    prev_was_blank = False
                    
                    if line_stripped.startswith('- ') or line_stripped.startswith('* '):
                        text = line_stripped[2:]
                        try:
                            p = doc.add_paragraph(style='List Bullet')
                        except KeyError:
                            p = doc.add_paragraph()
                            run = p.add_run('\u2022 ')
                            run.font.name = 'Symbol'
                        _add_formatted_runs(p, text)
                        continue
                    
                    list_match = re.match(r'^(\d+)[.)]\s+(.*)', line_stripped)
                    if list_match:
                        text = list_match.group(2)
                        try:
                            p = doc.add_paragraph(style='List Number')
                        except KeyError:
                            p = doc.add_paragraph()
                            p.add_run(f'{list_match.group(1)}. ')
                        _add_formatted_runs(p, text)
                        continue
                    
                    heading_match = re.match(r'^(#{1,4})\s+(.*)', line_stripped)
                    if heading_match:
                        level = min(len(heading_match.group(1)), 4)
                        text = heading_match.group(2)
                        try:
                            doc.add_heading(text, level=level + 1)
                        except Exception:
                            doc.add_paragraph(text)
                        continue
                    
                    p = doc.add_paragraph()
                    _add_formatted_runs(p, line_stripped)
            
            # Add images
            for img in images:
                img_path = img.get("path", "")
                img_caption = img.get("caption", "")
                img_width_inches = img.get("width", 5.0)
                
                if not img_path:
                    continue
                
                try:
                    img_file = resolve_workspace_path(session_id, img_path)
                    if img_file.exists():
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(str(img_file), width=Inches(img_width_inches))
                        
                        if img_caption:
                            cap_p = doc.add_paragraph()
                            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cap_run = cap_p.add_run(img_caption)
                            cap_run.font.size = Pt(9)
                            cap_run.font.italic = True
                except Exception as e:
                    # Add placeholder if image fails
                    doc.add_paragraph(f"[Image: {img_path} — failed to insert: {e}]")
        
        doc.save(str(out_file))
        _chown_alloc(out_file, session_id)

        result_size = out_file.stat().st_size
        
        # Build a summary of what was done
        section_names = [s.get("heading", f"Section {i+1}") for i, s in enumerate(sections)]
        total_images = sum(len(s.get("images", [])) for s in sections)
        
        return make_success({
            "output_path": str(out_file),
            "size": result_size,
            "sections_added": len(sections),
            "section_names": section_names,
            "images_inserted": total_images,
            "cover_preserved": keep_cover_page,
            "cover_replacements_applied": len(cover_replacements),
            "summary": f"Created {output_path} from template {template_path}: {len(sections)} sections, {total_images} images, cover {'preserved' if keep_cover_page else 'not preserved'}",
        })
    
    except Exception as e:
        import traceback
        return make_error(f"Template fill failed: {traceback.format_exc()}", 500)


def _add_formatted_runs(paragraph, text: str):
    """Add runs with **bold** and *italic* formatting to a paragraph."""
    import re
    from docx.shared import Pt
    
    pattern = r'(\*\*.*?\*\*|\*.*?\*)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX Build (C# + OpenXML SDK creation route)
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/docx/build", methods=["POST"])
def docx_build() -> Response:
    """Build a DOCX file using the docx skill's C# + OpenXML SDK pipeline."""
    data = request.get_json(force=True) or {}
    session_id = data.get("session_id", "default")
    output_path = data.get("output_path", "")
    program_cs = data.get("program_cs", "")

    if not output_path:
        return make_error("Missing 'output_path' field", 400)

    session_workspace = ensure_session_dirs(session_id)

    try:
        out_file = resolve_workspace_path(session_id, output_path)
    except ValueError:
        return make_error("Invalid output_path", 400)

    # Surface-level setup (mkdir of parent + work dir + template copies) runs
    # under the session's alloc uid so the dotnet subprocess (which drops to the
    # same uid) can read/write these files.
    with as_session_uid(session_id):
        out_file.parent.mkdir(parents=True, exist_ok=True)
        work_dir = session_workspace / ".docx-work"
        work_dir.mkdir(parents=True, exist_ok=True)

        if program_cs:
            skill_dir = SKILLS_ROOT / "docx"
            csproj_src = skill_dir / "assets" / "templates" / "Docx.csproj"
            program_src = skill_dir / "assets" / "templates" / "Program.cs"
            csproj_dst = work_dir / "Docx.csproj"
            program_dst = work_dir / "Program.cs"
            if csproj_src.exists() and not csproj_dst.exists():
                csproj_dst.write_text(csproj_src.read_text(encoding="utf-8"), encoding="utf-8")
            if program_src.exists() and not program_dst.exists():
                program_dst.write_text(program_src.read_text(encoding="utf-8"), encoding="utf-8")
            program_dst.write_text(program_cs, encoding="utf-8")

    # Determine absolute output path for the script
    abs_output = str(out_file)

    docx_script = SKILLS_ROOT / "docx" / "scripts" / "docx"
    if not docx_script.exists():
        return make_error("docx skill script not found at /app/skills/docx/scripts/docx", 500)

    cmd = ["bash", str(docx_script), "build", abs_output]

    env = _alloc_run_env(session_id)
    env["DOCX_WORK_DIR"] = str(work_dir)

    start = time.time()
    try:
        jr = _run_as_alloc(session_id, cmd, env, str(work_dir), 300)
        duration_ms = int((time.time() - start) * 1000)

        if jr.get("timed_out"):
            return _capture_failure("DOCX build timed out", jr)
        if jr["exit_code"] != 0:
            return _capture_failure("DOCX build failed", jr)

        with as_session_uid(session_id):
            if not out_file.exists():
                return make_error("DOCX was not generated", 500)
            size = out_file.stat().st_size

        return make_success({
            "output_path": str(out_file),
            "size": size,
            "duration_ms": duration_ms,
            "stdout": jr["stdout"][-MAX_OUTPUT_SIZE:] if len(jr["stdout"]) > MAX_OUTPUT_SIZE else jr["stdout"],
        })
    except subprocess.TimeoutExpired:
        return make_error("DOCX build timed out", 504)
    except Exception as e:
        return make_error(f"DOCX build error: {e}", 500)


# ═══════════════════════════════════════════════════════════════════════════════
# PPTX (PPTD skill) — runs the bundled kimi_pptd binary
# ═══════════════════════════════════════════════════════════════════════════════

# The kimi_pptd runtime ships as a Nuitka-compiled ELF plus sibling .so libs
# under skills/pptx/scripts/runtime/. The skills mount is read-only, and the
# host may not preserve the executable bit (Windows copies, some CI checkouts).
# We therefore use the in-place binary when it is already executable, and fall
# back to a cached, chmod'd copy in /tmp otherwise.
_KIMI_RUNTIME_CACHE = Path("/tmp/kimi_pptd_runtime")
_KIMI_RUNTIME_LOCK = Path("/tmp/kimi_pptd_runtime.lock")

# Packages whose data files must exist on disk next to the binary even though
# their code is bytecode-compiled into kimi_pptd. certifi ships cacert.pem
# (the default TLS trust bundle that `requests` loads at import time); without
# it, `import requests` raises FileNotFoundError in adapters.py and the whole
# pptd->pptx convert path crashes.
def _ensure_runtime_data_pkgs(target_dir: Path) -> None:
    try:
        import importlib
        for pkg in ("certifi",):
            mod = importlib.import_module(pkg)
            mod_dir = Path(getattr(mod, "__file__", "")).parent
            if not mod_dir.name == pkg or not mod_dir.is_dir():
                continue
            dst = target_dir / pkg
            if dst.is_dir():
                # Refresh only if cacert.pem is missing or stale.
                ca_src = mod_dir / "cacert.pem"
                ca_dst = dst / "cacert.pem"
                if ca_dst.exists():
                    continue
            if dst.exists():
                shutil.rmtree(dst)
            shutil.copytree(mod_dir, dst)
            # The global umask is 077; copytree's own child dirs may end up at
            # 0700, which would block the dropped alloc uid from traversing the
            # copied package tree. Re-open the whole certifi subtree.
            for root, _dirs, files in os.walk(dst):
                os.chmod(root, 0o755)
                for name in files:
                    os.chmod(os.path.join(root, name), 0o644)
        # cacert.pem sanity check — if still missing, copy it from certifi.
        ca = target_dir / "certifi" / "cacert.pem"
        if not ca.exists():
            try:
                import certifi
                ca.parent.mkdir(parents=True, exist_ok=True)
                shutil.copyfile(certifi.where(), ca)
                os.chmod(ca, 0o644)
            except Exception:
                pass
    except Exception as e:  # pragma: no cover
        logger.warning("kimi_pptd runtime data-package sync failed: %s", e)


def _kimi_cache_ready(cache_binary: Path) -> bool:
    """True only when the cached binary exists AND is executable by the
    non-root alloc child (S_IXOTH). Verifying the actual exec bit — instead of a
    `.ready` sentinel — is what makes the cache self-healing: a partially
    prepared or wrong-permission cache is detected and rebuilt rather than
    silently causing ``[Errno 13] Permission denied`` in the dropped child."""
    try:
        if not cache_binary.exists():
            return False
        return bool(cache_binary.stat().st_mode & 0o001)  # S_IXOTH
    except OSError:
        return False


def _prepare_kimi_cache(runtime_dir: Path) -> Path:
    """Copy the read-only kimi_pptd runtime into a writable, exec-permitted cache
    under ``/tmp/kimi_pptd_runtime``, returning the cached binary path. Safe
    across gunicorn sync workers: preparation is serialized with a file lock so
    two workers never race on ``shutil.copytree`` of the same destination (the
    cause of intermittent permission/corruption errors)."""
    _KIMI_RUNTIME_CACHE.parent.mkdir(parents=True, exist_ok=True)
    cache_binary = _KIMI_RUNTIME_CACHE / "kimi_pptd"

    # The lock file lives as a sibling in /tmp (mode 1777), so every worker can
    # create/open it; holding LOCK_EX serializes all preparation. The file is
    # intentionally persistent and reused across requests/workers — unlinking
    # it would introduce a TOCTOU race (two workers could each create+open a
    # different inode and both think they hold the lock). It holds no data
    # (advisory flock only), costs zero bytes, and is created once per a
    # container's lifetime. Created 0644 (root-owned) so it's inspectable for
    # debugging rather than the umask-077 default of 0600.
    lf = open(_KIMI_RUNTIME_LOCK, "a+")
    try:
        st = os.fstat(lf.fileno())
        if st.st_mode & 0o777 != 0o644:
            os.chmod(_KIMI_RUNTIME_LOCK, 0o644)
    except OSError:
        pass
    try:
        fcntl.flock(lf, fcntl.LOCK_EX)
        # Re-check under the lock: another worker may have prepared it while we
        # waited.
        if _kimi_cache_ready(cache_binary):
            _ensure_runtime_data_pkgs(_KIMI_RUNTIME_CACHE)
            return cache_binary

        # (Re)build from scratch. rmtree is safe here because we hold the lock.
        if _KIMI_RUNTIME_CACHE.exists():
            shutil.rmtree(_KIMI_RUNTIME_CACHE, ignore_errors=True)
        shutil.copytree(runtime_dir, _KIMI_RUNTIME_CACHE)

        # kimi_pptd writes oxml template fragments next to itself at runtime;
        # ensure the dir exists and is traversable.
        (_KIMI_RUNTIME_CACHE / "pptx" / "oxml").mkdir(parents=True, exist_ok=True)

        # Force the whole cache tree to be traversable+readable by the non-root
        # alloc child. Directories get 0o755 (o+x is required for traversal);
        # the binary and the shared libraries (.so / .so.*) get 0o755 so the
        # kernel can exec / mmap them; everything else stays 0o644. The global
        # ``umask 077`` would otherwise leave these at 0700/0600 and the child
        # would get EACCES — exactly the intermittent failure this fixes.
        for root, dirs, files in os.walk(_KIMI_RUNTIME_CACHE):
            os.chmod(root, 0o755)
            for name in dirs:
                os.chmod(os.path.join(root, name), 0o755)
            for name in files:
                fp = os.path.join(root, name)
                if name == "kimi_pptd" or name.endswith(".so") or ".so." in name:
                    os.chmod(fp, 0o755)
                else:
                    os.chmod(fp, 0o644)
        os.chmod(_KIMI_RUNTIME_CACHE, 0o755)
        os.chmod(cache_binary, 0o755)

        _ensure_runtime_data_pkgs(_KIMI_RUNTIME_CACHE)

        if not _kimi_cache_ready(cache_binary):
            raise RuntimeError(
                "kimi_pptd cache binary is not executable after setup "
                f"(mode {oct(cache_binary.stat().st_mode)})"
            )
        return cache_binary
    finally:
        try:
            fcntl.flock(lf, fcntl.LOCK_UN)
        finally:
            lf.close()


def _resolve_kimi_pptd() -> tuple[Path, Path]:
    """Return (binary_path, runtime_dir) for the kimi_pptd executable.

    Prefers the mounted binary when it is already executable by others (fast
    path, no copy). Otherwise prepares (and validates) the writable cache copy.
    The runtime_dir returned is the directory whose libs/templates the binary
    expects to find next to itself (used for LD_LIBRARY_PATH)."""
    runtime_dir = SKILLS_ROOT / "pptx" / "scripts" / "runtime"
    binary = runtime_dir / "kimi_pptd"
    if not binary.exists():
        raise FileNotFoundError(
            "kimi_pptd binary not found at /app/skills/pptx/scripts/runtime/kimi_pptd"
        )

    # Fast path: mounted binary is executable BY OTHERS (the alloc child runs as
    # a non-root uid, so it needs the other-x bit).
    try:
        if binary.stat().st_mode & 0o001:  # S_IXOTH
            return binary, runtime_dir
    except OSError:
        pass

    return _prepare_kimi_cache(runtime_dir), _KIMI_RUNTIME_CACHE


def _build_pptx_result(
    session_id: str,
    action: str,
    in_file: Path,
    out_file: Optional[Path],
    jr: dict,
    duration_ms: int,
) -> dict[str, Any]:
    """Assemble the structured kimi_pptd result record from the raw run journal.

    A non-zero exit code (or a timeout) is NOT turned into an HTTP error here:
    the caller (the orchestrator) reads ``exit_code``/``timed_out`` from the
    structured result, and — crucially — the ``stdout``/``stderr`` stay attached
    so the UI shows the checker/render logs instead of an opaque failure.
    """
    stdout = jr.get("stdout", "") or ""
    stderr = jr.get("stderr", "") or ""
    result: dict[str, Any] = {
        "action": action,
        "exit_code": jr["exit_code"],
        "stdout": stdout[-MAX_OUTPUT_SIZE:] if len(stdout) > MAX_OUTPUT_SIZE else stdout,
        "stderr": stderr[-MAX_OUTPUT_SIZE:] if len(stderr) > MAX_OUTPUT_SIZE else stderr,
        "duration_ms": duration_ms,
    }
    if jr.get("timed_out"):
        result["timed_out"] = True
        result["error"] = jr.get("error") or f"kimi_pptd {action} timed out"

    session_workspace = (WORKSPACE_ROOT / session_id).resolve()
    if out_file is not None:
        with as_session_uid(session_id):
            if not out_file.exists():
                out_exists = False
                out_is_file = False
                out_size = 0
                rel_files: list[str] = []
            else:
                out_exists = True
                out_is_file = out_file.is_file()
                out_size = out_file.stat().st_size if out_is_file else 0
                rel_files = []
                if action == "screenshot":
                    for child in sorted(out_file.iterdir()):
                        if child.is_file() and child.suffix.lower() in (".png", ".jpg", ".jpeg", ".webp"):
                            try:
                                rel_files.append(str(child.relative_to(session_workspace)))
                            except ValueError:
                                rel_files.append(child.name)
        if out_exists:
            if out_is_file:
                result["output_path"] = str(out_file)
                result["size"] = out_size
            elif action == "screenshot":
                result["output_dir"] = str(out_file)
                result["images"] = rel_files
            else:
                result["output_dir"] = str(out_file)
    return result


def _prepare_pptx_run(
    session_id: str,
    action: str,
    input_path: str,
    output_path: str,
    pages: str,
) -> tuple[Path, Optional[Path], Path, Path, list[str], str]:
    """Validate + resolve paths and build the argv/env for a kimi_pptd run.

    Returns ``(in_file, out_file, binary, runtime_dir, argv_tail, work_dir)``.
    Raises ``ValueError`` for client-side input errors (bad paths), and bubbles
    up ``FileNotFoundError`` if the kimi_pptd binary is unavailable. Split out so
    both the JSON and streaming routes share identical preparation."""
    if action not in ("check", "convert", "screenshot"):
        raise ValueError("Invalid 'action'; must be one of check, convert, screenshot")
    if not input_path:
        raise ValueError("Missing 'input_path'")

    ensure_session_dirs(session_id)
    in_file = resolve_workspace_path(session_id, input_path)

    out_file: Optional[Path] = None
    with as_session_uid(session_id):
        if not in_file.exists():
            raise FileNotFoundError(f"Input file not found: {input_path}")
        if output_path:
            out_file = resolve_workspace_path(session_id, output_path)
            out_file.parent.mkdir(parents=True, exist_ok=True)
            if action == "screenshot" or output_path.endswith("/"):
                out_file.mkdir(parents=True, exist_ok=True)

    binary, runtime_dir = _resolve_kimi_pptd()

    argv_tail: list[str] = []
    if action in ("convert", "screenshot") and out_file is not None:
        argv_tail += ["-o", str(out_file)]
    if action == "screenshot" and pages:
        argv_tail += ["-p", str(pages)]

    env = _alloc_run_env(session_id, base=os.environ.copy())
    env["PYTHONUTF8"] = "1"
    env.setdefault("LC_ALL", "C.UTF-8")
    env.setdefault("LANG", "C.UTF-8")
    env["LD_LIBRARY_PATH"] = (
        str(runtime_dir) + os.pathsep + env.get("LD_LIBRARY_PATH", "")
    )

    return in_file, out_file, binary, runtime_dir, argv_tail, env


@app.route("/pptx/run", methods=["POST"])
def pptx_run() -> Response:
    """Run a kimi_pptd subcommand (check | convert | screenshot) on a
    workspace file. This is the PPTD skill's execution backend."""
    data = request.get_json(force=True) or {}
    session_id = data.get("session_id", "default")
    action = (data.get("action") or "").strip()
    input_path = data.get("input_path", "")
    output_path = data.get("output_path", "")
    pages = data.get("pages", "")

    try:
        in_file, _out_file, binary, _runtime_dir, argv_tail, env = _prepare_pptx_run(
            session_id, action, input_path, output_path, pages
        )
    except ValueError as e:
        return make_error(str(e), 400)
    except FileNotFoundError as e:
        status = 404 if str(e).startswith("Input file not found") else 500
        return make_error(str(e), status)

    argv = [str(binary), action, str(in_file)] + argv_tail
    start = time.time()
    try:
        jr = _run_as_alloc(session_id, argv, env, str(in_file.parent), 300)
        duration_ms = int((time.time() - start) * 1000)
        result = _build_pptx_result(session_id, action, in_file, _out_file, jr, duration_ms)
        # Non-zero exit / timeout are returned as a structured success (HTTP 200)
        # so the orchestrator keeps stdout/stderr and surfaces them in the UI
        # instead of losing them in an error envelope.
        return make_success(result)
    except subprocess.TimeoutExpired:
        result = {
            "action": action, "exit_code": -1, "stdout": "", "stderr": "",
            "duration_ms": int((time.time() - start) * 1000),
            "timed_out": True, "error": f"kimi_pptd {action} timed out",
        }
        return make_success(result)
    except Exception as e:  # pragma: no cover
        logger.exception("pptx_run error")
        return make_error(f"kimi_pptd {action} error: {e}", 500)


@app.route("/pptx/run/stream", methods=["POST"])
def pptx_run_stream():
    """Stream kimi_pptd output live as newline-delimited JSON, matching the
    ``/exec/python/stream`` protocol so the orchestrator can attach the same
    chunk-forwarding callback:

      {"t":"stdout","s":"..."} / {"t":"stderr","s":"..."}  — live chunks
      {"t":"result", ...}                                   — final result record

    A non-zero exit / timeout is delivered as a ``result`` record (never an HTTP
    error) so the streamed logs are not truncated before the client sees them.
    """
    data = request.get_json(force=True) or {}
    session_id = data.get("session_id", "default")
    action = (data.get("action") or "").strip()
    input_path = data.get("input_path", "")
    output_path = data.get("output_path", "")
    pages = data.get("pages", "")

    try:
        in_file, _out_file, binary, _runtime_dir, argv_tail, env = _prepare_pptx_run(
            session_id, action, input_path, output_path, pages
        )
    except ValueError as e:
        return make_error(str(e), 400)
    except FileNotFoundError as e:
        status = 404 if str(e).startswith("Input file not found") else 500
        return make_error(str(e), status)

    argv = [str(binary), action, str(in_file)] + argv_tail

    def generate():
        import queue as _q
        import threading as _t

        out_q: "_q.Queue" = _q.Queue()

        def worker():
            start = time.time()
            try:
                def cb(stream, text):
                    out_q.put({"t": stream, "s": text})
                jr = _run_as_alloc(session_id, argv, env, str(in_file.parent), 300, on_chunk=cb)
                duration_ms = int((time.time() - start) * 1000)
                result = _build_pptx_result(session_id, action, in_file, _out_file, jr, duration_ms)
                out_q.put({"t": "result", "data": result})
            except subprocess.TimeoutExpired:
                out_q.put({"t": "result", "data": {
                    "action": action, "exit_code": -1, "stdout": "", "stderr": "",
                    "duration_ms": int((time.time() - start) * 1000),
                    "timed_out": True, "error": f"kimi_pptd {action} timed out",
                }})
            except Exception as e:  # pragma: no cover
                logger.exception("pptx_run_stream error")
                out_q.put({"t": "result", "data": {
                    "action": action, "exit_code": -1, "stdout": "", "stderr": "",
                    "duration_ms": int((time.time() - start) * 1000),
                    "error": f"kimi_pptd {action} error: {e}",
                }})

        th = _t.Thread(target=worker, daemon=True)
        th.start()
        try:
            while True:
                item = out_q.get(timeout=330)
                yield json.dumps(item) + "\n"
                if item.get("t") == "result":
                    break
        finally:
            th.join(timeout=0)

    return Response(generate(), mimetype="application/x-ndjson")


# ═══════════════════════════════════════════════════════════════════════════════
# Conversions
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/convert/html-to-pdf", methods=["POST"])
def convert_html_to_pdf() -> Response:
    data = request.get_json(force=True) or {}
    input_path = data.get("input_path", "")
    output_path = data.get("output_path", "")
    options = data.get("options", {})
    session_id = data.get("session_id", "default")

    if not input_path or not output_path:
        return make_error("Missing input_path or output_path", 400)

    try:
        in_file = resolve_workspace_path(session_id, input_path)
        out_file = resolve_workspace_path(session_id, output_path)
    except ValueError:
        return make_error("Invalid path", 400)

    with as_session_uid(session_id):
        if not in_file.exists():
            return make_error("Input HTML file not found", 404)
        out_file.parent.mkdir(parents=True, exist_ok=True)

    # Use the Python Playwright-based converter. The image ships only the
    # Python `playwright` package (pip) + its browsers at
    # $PLAYWRIGHT_BROWSERS_PATH; the Node `playwright` npm module is not
    # installed, so the old node html_to_pdf.js failed with MODULE_NOT_FOUND.
    script_path = Path("/app/scripts/html_to_pdf.py")
    if not script_path.exists():
        return make_error("html_to_pdf.py script not found", 500)

    cmd = [
        sys.executable, str(script_path),
        str(in_file), str(out_file),
        json.dumps(options),
    ]

    start = time.time()
    try:
        env = _alloc_run_env(session_id)
        jr = _run_as_alloc(session_id, cmd, env, str(in_file.parent), 120)
        duration_ms = int((time.time() - start) * 1000)

        if jr.get("timed_out"):
            return _capture_failure("PDF conversion timed out", jr)
        if jr["exit_code"] != 0:
            return _capture_failure("PDF conversion failed", jr)

        with as_session_uid(session_id):
            if not out_file.exists():
                return make_error("PDF was not generated", 500)
            size = out_file.stat().st_size

        return make_success({
            "output_path": str(out_file),
            "size": size,
            "duration_ms": duration_ms,
        })
    except subprocess.TimeoutExpired:
        return make_error("PDF conversion timed out", 504)
    except Exception as e:
        return make_error(f"PDF conversion error: {e}", 500)


@app.route("/convert/docx-to-pdf", methods=["POST"])
def convert_docx_to_pdf() -> Response:
    data = request.get_json(force=True) or {}
    input_path = data.get("input_path", "")
    output_path = data.get("output_path", "")
    session_id = data.get("session_id", "default")

    if not input_path or not output_path:
        return make_error("Missing input_path or output_path", 400)

    try:
        in_file = resolve_workspace_path(session_id, input_path)
        out_file = resolve_workspace_path(session_id, output_path)
    except ValueError:
        return make_error("Invalid path", 400)

    with as_session_uid(session_id):
        if not in_file.exists():
            return make_error("Input DOCX file not found", 404)
        out_file.parent.mkdir(parents=True, exist_ok=True)

    # Use LibreOffice headless conversion. HOME points at a per-session
    # alloc-owned .home so LibreOffice/dconf can write their profile caches
    # (shared /tmp was the cause of the dconf "Permission denied" fatal error).
    # -env:UserInstallation forces LO's own user profile into that home too, so
    # "User installation could not be completed" cannot recur even if some
    # downstream tool resets HOME.
    out_dir = out_file.parent
    home_uri = "file://" + str(_session_home(session_id))
    cmd = [
        "libreoffice",
        "--headless",
        "--nologo",
        "-env:UserInstallation=" + home_uri,
        "--convert-to", "pdf",
        "--outdir", str(out_dir),
        str(in_file),
    ]
    env = _alloc_run_env(session_id)

    start = time.time()
    try:
        jr = _run_as_alloc(session_id, cmd, env, str(in_file.parent), 120)
        duration_ms = int((time.time() - start) * 1000)

        with as_session_uid(session_id):
            expected_output = out_dir / (in_file.stem + ".pdf")
            if expected_output.exists() and expected_output != out_file:
                expected_output.rename(out_file)

        if jr.get("timed_out"):
            return _capture_failure("DOCX to PDF conversion timed out", jr)
        with as_session_uid(session_id):
            if not out_file.exists():
                return _capture_failure("LibreOffice conversion failed", jr)
            size = out_file.stat().st_size

        return make_success({
            "output_path": str(out_file),
            "size": size,
            "duration_ms": duration_ms,
        })
    except subprocess.TimeoutExpired:
        return make_error("DOCX to PDF conversion timed out", 504)
    except Exception as e:
        return make_error(f"Conversion error: {e}", 500)


# ═══════════════════════════════════════════════════════════════════════════════
# OCR (PaddleOCR-VL via llama.cpp)
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/ocr/status", methods=["GET"])
def ocr_status() -> Response:
    """Report OCR engine readiness (active = models downloaded + server live).

    The app queries this to decide whether to expose the OCR tool to the agent,
    so it must be cheap and never crash even if the engine was never set up.
    """
    try:
        return jsonify(ocr_engine.get_status())
    except Exception as e:  # pragma: no cover
        logger.exception("ocr_status error")
        return jsonify({
            "active": False, "ready": False, "state": "unknown",
            "message": f"status check failed: {e}", "errors": [str(e)],
        })


@app.route("/ocr", methods=["POST"])
def ocr_run() -> Response:
    """Run an OCR task on a workspace image or PDF.

    Body: { session_id, input_path, task } where task is one of
    ocr | table | chart | formula | spotting | seal. Runs as root (it must read
    session files and talk to the shared root-owned llama-server); the extracted
    text is returned to the agent.
    """
    data = request.get_json(force=True) or {}
    session_id = data.get("session_id", "default")
    input_path = (data.get("input_path") or "").strip()
    task = (data.get("task") or "ocr").strip()

    if task not in ocr_engine.VALID_TASKS:
        return make_error(
            f"Invalid 'task'; must be one of {', '.join(ocr_engine.VALID_TASKS)}", 400
        )
    if not input_path:
        return make_error("Missing 'input_path'", 400)

    ensure_session_dirs(session_id)
    try:
        in_file = resolve_workspace_path(session_id, input_path)
    except ValueError:
        return make_error("Invalid input_path", 400)
    if not in_file.exists():
        return make_error(f"Input file not found: {input_path}", 404)

    start = time.time()
    try:
        res = ocr_engine.handle_ocr(str(in_file), task)
        res["duration_ms"] = int((time.time() - start) * 1000)
        return make_success(res)
    except ocr_engine.OcrUnavailable as e:
        # 503 so the app can surface "OCR not ready" distinctly from a bad request.
        return make_error(str(e), 503)
    except Exception as e:  # pragma: no cover
        logger.exception("ocr_run error")
        return make_error(f"OCR error: {e}", 500)


# ═══════════════════════════════════════════════════════════════════════════════
# Main
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Chatinterface Agent Sandbox API")
    parser.add_argument("--port", type=int, default=8080)
    parser.add_argument("--host", type=str, default="0.0.0.0")
    args = parser.parse_args()

    logger.info(f"Starting sandbox server on {args.host}:{args.port}")
    # threaded=False is REQUIRED for per-session seteuid security: only one
    # request may manipulate the effective uid of this process at a time. Use
    # gunicorn --worker-class sync for production.
    app.run(host=args.host, port=args.port, threaded=False)
